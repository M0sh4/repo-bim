import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement, ns


class ReportGenerator:
    def __init__(self):
        self.document = Document()
        self.num_titles = 0
        self.num_titles_sub = 0
        self.document = Document()
        self.document.styles['Normal'].font.name = "Verdana"  # Aplica Verdana a todo el documento
        self.document.styles['Normal'].font.size = Pt(10)  # Tamaño estándar

        self.style_bold_red_left = {
            "bold": True,
            "background_color": RGBColor(255, 0, 0),
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "text_color": RGBColor(255, 255, 255),
            "vertical_alignment": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "font_name": "Verdana",  # Añadimos la fuente Verdana
            "font_size": Pt(10)
        }

        self.style_bold_red_center = {
            "bold": True,
            "background_color": RGBColor(255, 0, 0),
            "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "text_color": RGBColor(255, 255, 255),
            "vertical_alignment": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "font_name": "Verdana",  # Añadimos la fuente Verdana
            "font_size": Pt(10)
        }

        self.style_red_text_left = {
            "bold": True,
            "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "text_color": RGBColor(255, 0, 0),
            "vertical_alignment": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "font_name": "Verdana",  # Añadimos la fuente Verdana
            "font_size": Pt(10)
        }

        self.style_black_enter = {
            "background_color": RGBColor(255, 255, 255),
            "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "text_color": RGBColor(0, 0, 0),
            "vertical_alignment": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            "font_name": "Verdana",  # Añadimos la fuente Verdana
            "font_size": Pt(10)
        }

    def add_header(self):
        logo_path = "automatizar_evidencias\\reports\\resources\\logo.jpg"
        seccion_header = self.document.sections[0].header
        datos = [
            ["", "INFORME DE PRUEBAS DEL SISTEMA"],
            ["", "Gerencia de la Tecnología, Innovación y Operaciones"],
            ["", "Desarrollo e Innovación"],
            ["", "Clasificación de la Información: Confidencial / Uso Interno"],
        ]

        format_matrix = [
            [None, self.style_bold_red_center],
            [None, self.style_red_text_left],
            [None, self.style_red_text_left],
            [None, None],
        ]

        tabla = seccion_header.add_table(rows=4, cols=2, width=Pt(560))
        tabla.style = "Table Grid"

        # Ajustar el ancho de las columnas
        self._set_column_width(tabla.columns[0], 3.34)
        self._set_column_width(tabla.columns[1], 15.27)

        for i, fila in enumerate(datos):
            celdas = tabla.rows[i].cells
            for j, valor in enumerate(fila):
                celdas[j].text = str(valor)

        self._set_style_table(tabla, format_matrix)

        # mergear celda logo
        celda_logo = tabla.cell(0, 0)
        celda_logo.merge(tabla.cell(3, 0))
        self._clean_cell_text(celda_logo)

        # Agregar la imagen al logo
        logo_celda = tabla.cell(0, 0)
        logo_parrafo = logo_celda.paragraphs[0]
        logo_run = logo_parrafo.add_run()
        logo_run.add_picture(logo_path, width=Mm(27), height=Mm(17))
        logo_parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _set_style_table(self, tabla, format_matrix):
        for i, row in enumerate(tabla.rows):
            for j, cell in enumerate(row.cells):
                style = format_matrix[i][j]
                if style:
                    if style.get("bold"):
                        self._set_cell_bold(cell)
                    if style.get("background_color"):
                        self._set_cell_background_color(cell, style["background_color"])
                    if style.get("alignment") is not None:
                        self._set_cell_alignment(cell, style["alignment"])
                    if style.get("text_color"):
                        self._set_cell_text_color(cell, style["text_color"])
                    if style.get("vertical_alignment"):
                        self._set_cell_vertical_alignment(cell, style["vertical_alignment"])
                else:
                    self._set_cell_alignment(cell, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def add_title(self, text, level=1):
        if level == 1:
            self.num_titles += 1
            self.num_titles_sub = 1
            self.num_titles_sub_sub = 1
            title_number_format = f"{self.num_titles}. {text}"
        elif level == 2:
            title_number_format = f"{self.num_titles}.{self.num_titles_sub} {text}"
            self.num_titles_sub += 1
            self.num_titles_sub_sub = 1
        elif level == 3:
            title_number_format = f"{self.num_titles}.{self.num_titles_sub - 1}.{self.num_titles_sub_sub} {text}"
            self.num_titles_sub_sub += 1
        else:
            raise ValueError("Unsupported title level. Only levels 1, 2, and 3 are supported.")

        title = self.document.add_heading(level=level)
        run = title.add_run(title_number_format)
        run.bold = True
        run.font.size = Pt(10)  # Ajustando el tamaño de la fuente
        run.font.name = "Verdana"  # Configurando la fuente a Arial
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if level == 1 else WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run.font.color.rgb = RGBColor(0, 0, 0)

    def add_paragraph(self, text, style=None, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        paragraph = self.document.add_paragraph(text)

        if style:
            paragraph.style = style

        paragraph.alignment = alignment

    def _set_cell_background_color(self, cell, color):
        cell_properties = cell._element.get_or_add_tcPr()  # pylint: disable=protected-access
        colorhex = str(color)
        shade = OxmlElement("w:shd")
        shade.set(ns.qn("w:val"), "clear")
        shade.set(ns.qn("w:color"), "auto")
        shade.set(ns.qn("w:fill"), colorhex)
        cell_properties.append(shade)

    def _set_cell_text_color(self, cell, color):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = color

    def _set_cell_bold(self, cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    def _set_cell_alignment(self, cell, alignment):
        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment

    def _set_cell_vertical_alignment(self, cell, vertical_alignment):
        cell.vertical_alignment = vertical_alignment

    def _set_column_width(self, column, width_cm):
        width_pt = width_cm * 28.35  # Convertir cm a puntos
        for cell in column.cells:
            cell.width = Pt(width_pt)

    def _create_table(self, datos):
        filas, columnas = len(datos), len(datos[0])
        tabla = self.document.add_table(rows=filas, cols=columnas)

        for i, fila in enumerate(datos):
            celdas = tabla.rows[i].cells
            for j, valor in enumerate(fila):
                celdas[j].text = str(valor)
        return tabla

    # pylint: disable=R0913
    def create_table_general_data(
        self,
        nombre_proyecto,
        etapa_del_proyecto,
        analista,
        sistema,
        fecha=None,
    ):
        if fecha is None:
            fecha = datetime.now().strftime("%d/%m/%Y")

        # 🔹 Agregar espacio antes de la tabla
        self.document.add_paragraph("")

        # 🔹 Definir datos para la tabla
        datos = [
            ["Nombre del Proyecto", nombre_proyecto],
            ["Etapa del Proyecto", etapa_del_proyecto],
            ["Analista QA", analista],
            ["Fecha de Informe", fecha],
            ["Nombre de Sistema", sistema]
        ]

        # 🔹 Crear la tabla de forma segura
        tabla = self._create_table(datos)
        self._style_general_data_table(tabla)

        # 🔹 Agregar espacio después de la tabla para separación
        self.document.add_paragraph("")


    def _style_general_data_table(self, tabla):
        """Aplica estilo y alineación correcta a la tabla de Datos Generales."""
        tabla.style = "Table Grid"

        # 🔹 Ajustar el ancho de columnas para una mejor distribución
        self._set_column_width(tabla.columns[0], 7.0)  # Columna de títulos (más ancha)
        self._set_column_width(tabla.columns[1], 13.0)  # Columna de valores

        # 🔹 Aplicar estilos celda por celda
        for i, row in enumerate(tabla.rows):
            for j, cell in enumerate(row.cells):
                self._set_cell_vertical_alignment(cell, WD_CELL_VERTICAL_ALIGNMENT.CENTER)

                if j == 0:  # Primera columna (títulos)
                    self._set_cell_background_color(cell, RGBColor(255, 0, 0))  # Rojo
                    self._set_cell_text_color(cell, RGBColor(255, 255, 255))  # Blanco
                    self._set_cell_bold(cell)
                    self._set_cell_alignment(cell, WD_PARAGRAPH_ALIGNMENT.LEFT)
                else:  # Segunda columna (valores)
                    self._set_cell_alignment(cell, WD_PARAGRAPH_ALIGNMENT.LEFT)

                # 🔹 Ajustar el espaciado dentro de la celda
                for paragraph in cell.paragraphs:
                    paragraph.space_before = Pt(4)  # Espacio antes del texto
                    paragraph.space_after = Pt(4)   # Espacio después del texto


    def create_table_stream_data(self, apk_name=None):
        if apk_name is None:
            apks_in_folder = self._get_files_in_folder("automatizar_evidencias/apk/android")
            apk_name = "\n".join(apks_in_folder)

        datos = [["Versión del APK"], [apk_name]]
        tabla = self._create_table(datos)
        self._style_stream_data_table(tabla)

    def _style_stream_data_table(self, tabla):
        tabla.style = "Table Grid"

        format_matrix = [[self.style_bold_red_center], [self.style_black_enter]]

        self._set_style_table(tabla, format_matrix)

    # pylint: disable=R0913
    def create_table_functional_test(
        self,
        caso_de_prueba,
        escenario_de_prueba,
        id_compartamos,
        tipo_billetera,
        numero_billetera_origen,
        numero_billetera_destino,
        id_operacion,
        definicion_caso,
        resultado_esperado,
        tipo_prueba,
        flujo_basico="<Flujo Básico>",
        log="<log>",
        fecha=None,
        hora=None
    ):
        if fecha is None:
            fecha = datetime.now().strftime("%d/%m/%Y")

        datos = [
            ["Caso de Prueba:", caso_de_prueba, "", ""],
            ["Escenario de Prueba:", escenario_de_prueba],
            ["ID Compartamos:", id_compartamos],
            ["Datos utilizados"],
            ["Tipo de Billetera:", tipo_billetera],
            ["Numero de Billetera de Origen:", numero_billetera_origen],
            ["Numero de Billetera de Destino:", numero_billetera_destino],
            ["ID de la Operación:", id_operacion],
            [
                "Definición del Caso de Prueba",
                "Resultado Esperado",
                "Tipo de Prueba",
                "Fecha y Hora de ejecución de la Prueba",
            ],
            [definicion_caso, resultado_esperado, tipo_prueba, fecha],
            ["Flujo"],
            [flujo_basico],
            ["Log"],
            [log],
        ]

        tabla = self._create_table(datos)
        self._style_table_functional_test(tabla)
        self.add_paragraph("")

    def _clean_cell_text(self, cell):
        original_alignment = cell.paragraphs[0].alignment
        text = cell.text.strip()
        lines = text.split("\n")
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        cleaned_text = " ".join(cleaned_lines)  # Join lines with a single space
        original_runs = cell.paragraphs[0].runs[:]  # Copy the original runs
        cell.text = ""  # Clear the cell text
        if cleaned_text:
            for run in original_runs:
                if run.text.strip():
                    new_run = cell.paragraphs[0].add_run(run.text.strip())
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.color.rgb = run.font.color.rgb
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name

        cell.paragraphs[0].alignment = original_alignment

    # pylint: disable=R0913
    def _merge_cells_in_range(self, table, start_row, end_row, start_col, end_col):
        for i in range(start_row, end_row + 1):
            row_cells = table.rows[i].cells
            for j in range(start_col, end_col):
                row_cells[start_col].merge(row_cells[j + 1])
                self._clean_cell_text(row_cells[start_col])

    def _style_table_functional_test(self, tabla):
        tabla.style = "Table Grid"
        self._set_column_width(tabla.columns[0], 4.5)

        format_matrix = [
            [self.style_bold_red_left, None, None, None],
            [self.style_bold_red_left, None, None, None],
            [self.style_bold_red_left, None, None, None],
            [self.style_bold_red_center, None, None, None],
            [self.style_bold_red_left, None, None, None],
            [self.style_bold_red_left, None, None, None],
            [self.style_bold_red_left, None, None, None],
            [self.style_bold_red_left, None, None, None],
            [
                self.style_bold_red_center,
                self.style_bold_red_center,
                self.style_bold_red_center,
                self.style_bold_red_center,
            ],
            [None, None, self.style_black_enter, self.style_black_enter],
            [self.style_bold_red_center, None, None, None],
            [self.style_black_enter, None, None, None],
            [self.style_bold_red_center, None, None, None],
            [self.style_black_enter, None, None, None],
        ]

        self._set_style_table(tabla, format_matrix)

        self._merge_cells_in_range(tabla, start_row=0, end_row=0, start_col=1, end_col=3)
        self._merge_cells_in_range(tabla, start_row=1, end_row=1, start_col=1, end_col=3)
        self._merge_cells_in_range(tabla, start_row=2, end_row=2, start_col=1, end_col=3)
        self._merge_cells_in_range(tabla, start_row=3, end_row=3, start_col=0, end_col=3)
        self._merge_cells_in_range(tabla, start_row=4, end_row=4, start_col=1, end_col=3)
        self._merge_cells_in_range(tabla, start_row=5, end_row=5, start_col=1, end_col=3)
        self._merge_cells_in_range(tabla, start_row=6, end_row=6, start_col=1, end_col=3)
        self._merge_cells_in_range(tabla, start_row=7, end_row=7, start_col=1, end_col=3)
        #self._merge_cells_in_range(tabla, start_row=8, end_row=8, start_col=0, end_col=3)
        self._merge_cells_in_range(tabla, start_row=10, end_row=10, start_col=0, end_col=3)
        self._merge_cells_in_range(tabla, start_row=11, end_row=11, start_col=0, end_col=3)
        self._merge_cells_in_range(tabla, start_row=12, end_row=12, start_col=0, end_col=3)
        self._merge_cells_in_range(tabla, start_row=13, end_row=13, start_col=0, end_col=3)

    def _get_files_in_folder(self, folder):
        if not os.path.isdir(folder):
            print(f"The folder '{folder}' does not exist.")
            return []

        files = os.listdir(folder)
        file_names = [file.split("\\")[-1] for file in files]
        return file_names

    def _add_text_and_images_in_cell(self, celda, contenido: list):
        for _, item in enumerate(contenido):
            if isinstance(item, str) and item.endswith((".jpg", ".png", ".jpeg", ".gif")):
                run = celda.add_paragraph().add_run()
                width_cm = 3
                width_pt = width_cm * 28.35  # Convertir cm a puntos
                run.add_picture(item, width=Pt(width_pt))
            else:
                p = celda.add_paragraph(item)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_content_in_cell(self, marcador, contenido: list):
        for tabla in self.document.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if marcador in celda.text:
                        celda.text = celda.text.replace(marcador, "")  # Borra el marcador del texto

                        current_paragraph = None  # Para agrupar imágenes en la misma línea
                        max_images_per_row = 4  # Número máximo de imágenes por fila antes de saltar a la siguiente

                        for item in contenido:
                            if isinstance(item, dict):
                                if item["tipo"] == "texto":
                                    # Añadir el título del paso como un párrafo independiente
                                    p = celda.add_paragraph(item["contenido"])
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                    current_paragraph = None  # Reiniciar para imágenes nuevas

                                elif item["tipo"] == "imagen":
                                    ruta_imagen = item["ruta"]
                                    if os.path.exists(ruta_imagen):
                                        if current_paragraph is None:
                                            current_paragraph = celda.add_paragraph()  # Nuevo párrafo para imágenes

                                        run = current_paragraph.add_run()
                                        
                                        # 🔥 Ajustar tamaño en píxeles (Convertido a puntos)
                                        ancho_px = 150  # Puedes ajustar este valor
                                        alto_px = 300  # Puedes ajustar este valor
                                        ancho_pt = ancho_px * 0.75  # Convertir px a puntos
                                        alto_pt = alto_px * 0.75
                                        
                                        run.add_picture(ruta_imagen, width=Pt(ancho_pt), height=Pt(alto_pt))

                                        # Alineación horizontal de imágenes
                                        if len(current_paragraph.runs) % max_images_per_row == 0:
                                            current_paragraph = None  # Nueva fila para imágenes



    def save_document(self, nombre_archivo):
        self.document.save(nombre_archivo)
