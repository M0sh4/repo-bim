import tkinter as tk
from tkinter import messagebox
from PIL import ImageGrab
from docx import Document
import os
import ctypes

# 🔹 Corregir problemas de escala en Windows
ctypes.windll.shcore.SetProcessDpiAwareness(1)  # Ajusta el DPI para capturas precisas

def capture_screen_region():
    """ Permite al usuario seleccionar una región exacta de la pantalla y la guarda en Word. """

    # Crear la ventana de selección
    root = tk.Tk()
    root.withdraw()  

    selection_window = tk.Toplevel(root)
    selection_window.overrideredirect(True)  
    selection_window.attributes("-topmost", True)  

    # Obtener dimensiones de pantalla
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    selection_window.geometry(f"{screen_width}x{screen_height}+0+0")

    # Fondo semi-transparente para resaltar la selección
    selection_window.configure(bg="black")
    selection_window.attributes("-alpha", 0.3)  

    # Crear lienzo para la selección con cursor de cruz
    canvas = tk.Canvas(selection_window, cursor="cross", bg="black", highlightthickness=0)
    canvas.pack(fill=tk.BOTH, expand=True)

    # Variables de selección
    start_x = start_y = None
    rect = None

    def on_button_press(event):
        """ Inicia la selección con clic. """
        nonlocal start_x, start_y, rect
        start_x, start_y = event.x, event.y
        rect = canvas.create_rectangle(start_x, start_y, start_x+1, start_y+1, outline="red", width=2)

    def on_drag(event):
        """ Expande la selección mientras se arrastra el mouse. """
        canvas.coords(rect, start_x, start_y, event.x, event.y)

    selection = {"x1": None, "y1": None, "x2": None, "y2": None}

    def on_button_release(event):
        """ Finaliza la selección y guarda coordenadas. """
        selection.update({
            "x1": min(start_x, event.x) + selection_window.winfo_rootx(),
            "y1": min(start_y, event.y) + selection_window.winfo_rooty(),
            "x2": max(start_x, event.x) + selection_window.winfo_rootx(),
            "y2": max(start_y, event.y) + selection_window.winfo_rooty()
        })
        selection_window.destroy()
        root.quit()

    # Enlazar eventos al lienzo
    canvas.bind("<Button-1>", on_button_press)
    canvas.bind("<B1-Motion>", on_drag)
    canvas.bind("<ButtonRelease-1>", on_button_release)

    root.mainloop()

    if None in selection.values() or selection["x1"] == selection["x2"] or selection["y1"] == selection["y2"]:
        print("⚠ Selección cancelada o inválida.")
        return None  

    # Capturar la región seleccionada
    return ImageGrab.grab(bbox=(selection["x1"], selection["y1"], selection["x2"], selection["y2"]))

def iniciar_captura():
    """ Ejecuta el flujo de captura de pantalla y guardado en Word. """
    imagen_capturada = capture_screen_region()

    if imagen_capturada:
        descripcion = descripcion_entry.get()

        # Crear documento Word y agregar descripción e imagen
        document = Document()
        document.add_paragraph(descripcion)

        # Guardar imagen y agregarla a Word
        output_folder = "automatizar_evidencias"
        os.makedirs(output_folder, exist_ok=True)
        image_path = os.path.join(output_folder, "captura_temp.png")

        imagen_capturada.save(image_path)
        document.add_picture(image_path)

        # Guardar documento
        word_path = os.path.join(output_folder, "Captura_Screen.docx")
        document.save(word_path)

        messagebox.showinfo("✅ Captura Guardada", f"Documento guardado: {word_path}")
    else:
        messagebox.showwarning("⚠ Error", "No se generó ninguna captura de pantalla.")

# --- 🌟 *Interfaz Principal* ---
ventana_principal = tk.Tk()
ventana_principal.title("Automatizar Evidencias")
ventana_principal.geometry("400x250")
ventana_principal.resizable(False, False)

# Etiqueta
label = tk.Label(ventana_principal, text="Ingrese una descripción para la captura:", font=("Arial", 12))
label.pack(pady=10)

# Campo de entrada
descripcion_entry = tk.Entry(ventana_principal, font=("Arial", 12), width=40)
descripcion_entry.pack(pady=5)

# Botón para iniciar la captura
boton_capturar = tk.Button(ventana_principal, text="📸 Capturar Pantalla", font=("Arial", 12, "bold"), command=iniciar_captura, bg="#28a745", fg="white")
boton_capturar.pack(pady=20)

# Iniciar la ventana principal
ventana_principal.mainloop()